from docx import Document
from docx.shared import Inches

class EarthquakeReport:
    def __init__(self, workingPath):
        self.workingPath = workingPath
        self.document = Document()

    def set_title(self, city, ruptureTime):
        self.document.add_heading("Earthquake Report for " + city + " on " + ruptureTime[:ruptureTime.find("_")], 0)

    def add_hazard_description(self, initialSummary, tectonicSummary):
        self.document.add_heading("Hazard Description", 1)

        self.document.add_paragraph(initialSummary)
        self.document.add_paragraph(tectonicSummary)

        self.document.add_picture(self.workingPath + "/overviewData/intensity.jpg", width=Inches(4.9), height=Inches(4.9))

    def add_sections(self, sectionsDictionary):
        for x in range(0, len(sectionsDictionary)):
            key = list(sectionsDictionary.items())[x][0]

            self.document.add_heading(key, 1)
            self.document.add_paragraph(sectionsDictionary[key])

    def add_picture(self, imagePath, width, height):
        self.document.add_picture(imagePath, width=Inches(width), height=Inches(height))

    def save(self):
        self.document.add_page_break()
        self.document.save(self.workingPath + "/briefing.docx")

class EarthquakeReportImage:
    def __init__(self, imagePath, width, height):
        self.imagePath = imagePath
        self.width = width
        self.height = height
